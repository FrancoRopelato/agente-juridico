import streamlit as st
import anthropic
import base64
import io
from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ─── CONFIGURACIÓN DE PÁGINA ───
st.set_page_config(
    page_title="Asistente Jurídico",
    page_icon="⚖️",
    layout="centered"
)

# ─── CSS PERSONALIZADO ───
st.markdown("""
<style>
    .app-header {
        background: #1a2744;
        padding: 16px 24px;
        border-radius: 10px 10px 0 0;
        display: flex;
        align-items: center;
        justify-content: space-between;
        margin-bottom: 0;
    }
    .app-header-left {
        display: flex;
        align-items: center;
        gap: 12px;
    }
    .app-logo {
        width: 36px;
        height: 36px;
        background: #c9a84c;
        border-radius: 6px;
        display: flex;
        align-items: center;
        justify-content: center;
        font-size: 18px;
    }
    .app-title {
        color: white;
        font-size: 15px;
        font-weight: 600;
        margin: 0;
    }
    .app-subtitle {
        color: rgba(255,255,255,0.5);
        font-size: 11px;
        margin: 0;
    }
    .app-user-badge {
        background: #1a2744;
        border: 1px solid rgba(255,255,255,0.3);
        border-radius: 4px;
        padding: 4px 10px;
        color: white;
        font-size: 10px;
        letter-spacing: 0.5px;
    }
    .stDownloadButton > button {
        background: #1a2744 !important;
        color: white !important;
        border: none !important;
        border-radius: 8px !important;
        padding: 10px 20px !important;
        font-weight: 500 !important;
        width: 100% !important;
    }
    .stButton > button {
        background: #1a2744 !important;
        color: white !important;
        border: none !important;
        border-radius: 8px !important;
        padding: 10px 20px !important;
        font-weight: 500 !important;
        width: 100% !important;
    }
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
</style>
""", unsafe_allow_html=True)

# ─── SYSTEM PROMPT ───
SYSTEM_PROMPT = """
IDENTIDAD Y ROL
Sos un asistente jurídico profesional especializado 
en derecho argentino. Tu función es asistir a los 
profesionales del estudio en el análisis de casos, 
elaboración de estrategias procesales y redacción 
de documentos judiciales y extrajudiciales.

No sos un abogado. No ejercés la profesión. Sos una 
herramienta de asistencia profesional. Toda decisión 
procesal y todo documento generado requiere revisión 
y aprobación del profesional autorizado antes de 
cualquier uso oficial.

Área de especialización activa: Derecho de Familia
Nombre del estudio: Estudio Cairoli

USUARIOS Y CONTEXTO
Los usuarios autorizados son profesionales del estudio.
Nunca interactuás directamente con clientes finales.
Si un mensaje sugiere que quien escribe es un cliente 
sin representación profesional, lo indicás claramente 
y no procedés hasta que un profesional confirme su presencia.

Adaptá el nivel de detalle según el perfil:
- Abogado titular o senior: comunicación directa, 
  sin explicar conceptos básicos, foco en estrategia.
- Abogado junior o pasante: incluí el razonamiento 
  detrás de cada decisión estratégica.

CAPACIDADES
Análisis de escritos judiciales y expedientes.
Elaboración de estrategias procesales y de negociación.
Redacción de escritos judiciales y extrajudiciales 
en formato procesal argentino listo para edición.
Resumen estructurado de expedientes.
Revisión jurídica de borradores.

LÍMITES
No afirmás resultados judiciales con certeza.
No presentás jurisprudencia sin clasificación de certeza.
No tomás decisiones procesales.
No respondés a clientes finales sin profesional presente.
No generás documentos oficiales sin recordatorio de revisión.
Advertís explícitamente cuando una consulta está fuera 
del área de especialización del estudio.

PROTOCOLO DE CITAS JURÍDICAS
Toda cita se clasifica obligatoriamente:
✓ CONFIRMADO: legislación vigente con artículo citado 
  o jurisprudencia de CSJN o Cámaras identificada.
⚠ PROBABLE: doctrina mayoritaria o jurisprudencia de 
  primera instancia. Verificar antes de citar.
✗ A VERIFICAR: requiere confirmación en base oficial 
  antes de cualquier uso.

FLUJO DE TRABAJO OBLIGATORIO

Paso 1 — Comprensión
Antes de cualquier análisis:
- Identificá el tipo de proceso o escrito solicitado
- Determiná si el estudio representa a actor o demandado
- Identificá todas las partes con su rol procesal exacto
- Revisá los documentos disponibles
- Si falta información crítica, solicitala antes de continuar

Paso 2 — Análisis del caso
Estructura obligatoria:
- Tipo de proceso
- Hechos relevantes identificados
- Puntos jurídicos controvertidos
- Argumentos de la parte contraria y sus debilidades
- Argumentos disponibles para la parte representada
- Prueba relevante identificada o a producir

ANÁLISIS DE CONTENIDO SUSTANCIAL — OBLIGATORIO
Antes de elaborar la estrategia, extraé y analizá 
explícitamente TODO lo que el documento acuerda, 
reclama o establece, incluyendo sin limitarse a:

ASPECTOS ECONÓMICOS:
- Montos reclamados o acordados
- Cuotas alimentarias — monto, periodicidad, 
  forma de pago, mecanismo de actualización
- Porcentajes aplicables
- Intereses y actualizaciones
- Valuaciones de bienes

ASPECTOS VINCULADOS A LOS HIJOS:
- Régimen de cuidado personal — con quién viven
- Régimen de comunicación y visitas — días, 
  horarios, modalidad
- Tiempos compartidos — vacaciones, feriados, 
  fechas especiales
- Responsabilidades de traslado — quién lleva 
  y busca, en qué horarios
- Gastos extraordinarios — salud, educación, 
  actividades extracurriculares
- Restricciones o condiciones especiales

ASPECTOS PATRIMONIALES:
- Bienes a dividir o ya divididos
- Atribución del hogar conyugal
- Deudas y responsabilidades compartidas
- Acuerdos sobre uso de bienes

OTROS ACUERDOS O RECLAMOS:
- Cualquier obligación de hacer o no hacer
- Plazos y condiciones establecidos
- Incumplimientos documentados
- Diferencias entre lo pactado anteriormente 
  y lo que se reclama ahora

RAZONABILIDAD Y ANÁLISIS CRÍTICO:
- ¿Lo acordado o reclamado es razonable según 
  jurisprudencia y práctica habitual?
- ¿Hay inconsistencias entre los hechos 
  alegados y lo que se pide?
- ¿Qué aspectos favorecen a cada parte?
- ¿Qué margen de negociación existe en cada punto?

Si alguno de estos aspectos no está en el documento, 
indicarlo explícitamente. No inventar información 
que no surge del texto.

Paso 3 — Estrategia procesal y de negociación
Generá tres caminos estratégicos alternativos:

CAMINO A — Estrategia agresiva
¿Cuál sería la posición más combativa posible?
¿Qué ventajas tiene? ¿Qué riesgos implica?

CAMINO B — Estrategia conservadora
¿Cuál sería la posición más segura y defensiva?
¿Qué sacrifica? ¿Qué protege?

CAMINO C — Estrategia de negociación
¿Cuál es el mejor escenario de acuerdo posible?
¿Qué necesita cada parte para cerrar?
¿Rangos razonables de acuerdo?

Recomendá el camino más adecuado para este caso 
con fundamento en los hechos y el perfil del cliente.

Paso 4 — Redacción del documento
- Estructura procesal argentina
- Todas las citas clasificadas según protocolo
- Formato listo para edición en Word
- Campos pendientes marcados con [COMPLETAR: descripción]

Paso 5 — Revisión jurídica
- Solidez de los argumentos jurídicos
- Coherencia entre hechos y derecho
- Adecuación de la estrategia procesal
- Posibles debilidades del planteo
- Sugerencias de mejora si las hay

ESCRITOS SUGERIDOS
Al finalizar el análisis, identificá el tipo de 
documento analizado y listá los escritos que 
corresponde redactar para este caso específico.
Nunca sugerás escritos que no correspondan al 
tipo de proceso o documento analizado.

Formato obligatorio de esta sección:
ESCRITOS QUE PODRÍA NECESITAR:
- [nombre del escrito]: [una línea explicando para qué sirve]

Si el documento no requiere ningún escrito de 
respuesta — indicarlo explícitamente y sugerí 
qué documentación complementaria podría ser útil.

RESUMEN DE EXPEDIENTE
Cuando el profesional proporcione un expediente 
completo para resumir, el resumen incluye:
- Carátula y número de expediente
- Partes intervinientes
- Tipo de proceso y fuero
- Hechos principales cronológicamente ordenados
- Pretensión del actor con fundamentos
- Argumentos jurídicos relevantes de cada parte
- Prueba ofrecida por cada parte
- Resoluciones importantes dictadas
- Estado actual del expediente
- Cuestiones controvertidas pendientes
- Próximos pasos procesales posibles

Extensión máxima una carilla. Ofrecer expansión 
de secciones si se requiere más detalle.

CIERRE OBLIGATORIO DE CADA DOCUMENTO
Todo documento cierra con:
─────────────────────────────
REVISIÓN PROFESIONAL REQUERIDA
Este documento requiere revisión y aprobación 
del abogado autorizado antes de uso oficial.

PRÓXIMOS PASOS SUGERIDOS: [acciones en orden]
DATOS PENDIENTES: [campos COMPLETAR del documento]
CITAS A VERIFICAR: [citas marcadas con ✗]
─────────────────────────────
"""

# ─── FUNCIONES ───
def procesar_pdf(archivo_bytes):
    """Convierte cualquier PDF a imágenes para Claude Vision."""
    from pdf2image import convert_from_bytes
    paginas = convert_from_bytes(archivo_bytes, dpi=150)
    imagenes = []
    for pagina in paginas:
        pagina.thumbnail((1500, 1500), Image.LANCZOS)
        buffer = io.BytesIO()
        pagina.save(buffer, format='JPEG', quality=85)
        imagen_b64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
        imagenes.append(imagen_b64)
    return imagenes

def obtener_analisis(imagenes, prompt_texto):
    """Manda las imágenes a Claude Vision y obtiene el análisis."""
    client = anthropic.Anthropic(api_key=st.secrets["ANTHROPIC_API_KEY"])
    
    mensaje_contenido = []
    for img in imagenes:
        mensaje_contenido.append({
            "type": "image",
            "source": {
                "type": "base64",
                "media_type": "image/jpeg",
                "data": img
            }
        })
    mensaje_contenido.append({
        "type": "text",
        "text": prompt_texto
    })
    
    message = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=4000,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": mensaje_contenido}]
    )
    return message.content[0].text

def generar_word(texto_analisis, nombre_documento):
    doc = Document()
    titulo = doc.add_heading('ANÁLISIS JURÍDICO', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo = doc.add_paragraph(f'Documento: {nombre_documento}')
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('─' * 60)
    lineas = texto_analisis.split('\n')
    for linea in lineas:
        linea = linea.strip()
        if not linea:
            doc.add_paragraph('')
            continue
        if linea.startswith('###') or linea.startswith('##'):
            doc.add_heading(linea.replace('###','').replace('##','').strip(), level=2)
        elif linea.startswith('#'):
            doc.add_heading(linea.replace('#','').strip(), level=1)
        elif linea.startswith('**') and linea.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(linea.replace('**',''))
            run.bold = True
        elif linea.startswith('- ') or linea.startswith('* '):
            doc.add_paragraph(linea[2:], style='List Bullet')
        elif linea.startswith('✓') or linea.startswith('⚠') or linea.startswith('✗'):
            doc.add_paragraph(linea, style='List Bullet')
        else:
            doc.add_paragraph(linea)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ─── HEADER ───
nombre_estudio = st.secrets.get("NOMBRE_ESTUDIO", "Estudio Jurídico")
nombre_usuario = st.secrets.get("NOMBRE_USUARIO", "Dr./Dra.")

st.markdown(f"""
<div class="app-header">
    <div class="app-header-left">
        <div class="app-logo">⚖️</div>
        <div>
            <p class="app-title">{nombre_estudio}</p>
            <p class="app-subtitle">Asistente jurídico con IA</p>
        </div>
    </div>
    <div class="app-user-badge">{nombre_usuario}</div>
</div>
""", unsafe_allow_html=True)

st.markdown("<div style='height: 16px'></div>", unsafe_allow_html=True)

# ─── UPLOAD ───
st.markdown("**Documento a analizar**")
archivo = st.file_uploader(
    "Subí el PDF del caso",
    type=['pdf'],
    help="Compatible con PDFs digitales y escaneados",
    label_visibility="collapsed"
)

if archivo:
    st.success(f"✓ {archivo.name}")
    
    if st.button("Analizar documento"):
        with st.spinner("Procesando documento con Vision... esto puede tomar unos segundos"):
            archivo_bytes = archivo.read()
            imagenes = procesar_pdf(archivo_bytes)
            st.info(f"📄 {len(imagenes)} página/s procesada/s")
            resultado = obtener_analisis(
                imagenes,
                "Analizá este documento jurídico y aplicá el flujo completo de trabajo obligatorio."
            )
            st.session_state['resultado'] = resultado
            st.session_state['imagenes'] = imagenes
            st.session_state['nombre_archivo'] = archivo.name

# ─── RESULTADO ───
if 'resultado' in st.session_state:
    resultado = st.session_state['resultado']
    
    st.markdown("---")
    st.markdown("🟢 **Análisis completado**")
    st.markdown(resultado)
    
    st.markdown("<div style='height: 8px'></div>", unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        word_buffer = generar_word(resultado, st.session_state['nombre_archivo'])
        st.download_button(
            label="⬇️ Descargar Word editable",
            data=word_buffer,
            file_name=f"analisis_{st.session_state['nombre_archivo'].replace('.pdf','')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    with col2:
        escrito_tipo = st.text_input(
            "¿Qué escrito necesitás redactar?",
            placeholder="Ej: convenio regulador, contestación de demanda, oficio...",
            label_visibility="visible"
        )
        
        if st.button("✍️ Redactar escrito") and escrito_tipo:
            with st.spinner(f"Redactando {escrito_tipo}..."):
                prompt_redaccion = f"""Basándote en el documento analizado, redactá un/a {escrito_tipo} completo en formato procesal argentino.
El escrito debe estar listo para editar y presentar.
Marcá con [COMPLETAR: descripción] los campos donde el abogado debe agregar datos específicos."""
                
                escrito = obtener_analisis(
                    st.session_state['imagenes'],
                    prompt_redaccion
                )
                st.session_state['escrito'] = escrito
                st.session_state['escrito_tipo'] = escrito_tipo

if 'escrito' in st.session_state:
    st.markdown("---")
    st.markdown(f"✍️ **{st.session_state['escrito_tipo'].capitalize()} redactado**")
    st.markdown(st.session_state['escrito'])
    
    word_escrito = generar_word(
        st.session_state['escrito'],
        f"{st.session_state['escrito_tipo']}_{st.session_state['nombre_archivo']}"
    )
    st.download_button(
        label=f"⬇️ Descargar {st.session_state['escrito_tipo']} en Word",
        data=word_escrito,
        file_name=f"{st.session_state['escrito_tipo']}_{st.session_state['nombre_archivo'].replace('.pdf','')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="dl_escrito"
    )
